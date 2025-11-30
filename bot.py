import telebot
from telebot import types
import pandas as pd
import json
from fpdf import FPDF # Оставлена, но не используется из-за проблем с кириллицей
from docx import Document
import re
import os

# --- 🔑 ВАШ АКТУАЛЬНЫЙ ТОКЕН 🔑 ---
BOT_TOKEN = "8413467526:AAHwY2YDRq5-bKvWW3FminyGmuPTZJ2eyVM" 
bot = telebot.TeleBot(BOT_TOKEN, threaded=False)

# --- ИНИЦИАЛИЗАЦИЯ ДАННЫХ И ФАЙЛОВ ---
DATA_FILE = "data.xlsx"
CATEGORIES_FILE = "categories.json"

# Убедимся, что DataFrame имеет нужные типы данных
initial_columns = ["user_id", "username", "category", "game", "votes", "poll_id"] 

if not os.path.exists(DATA_FILE):
    df = pd.DataFrame(columns=initial_columns)
    df.to_excel(DATA_FILE, index=False)
    df = pd.read_excel(DATA_FILE) # Перезагружаем для использования

if not os.path.exists(CATEGORIES_FILE):
    categories = [
        "Лучшая игра 2025 года",
        "Лучшее РПГ 2025 года",
        "Лучшая Action-Adventure 2025 года",
        "Лучшее DLC 2025 года",
        "Лучший шутер 2025 года",
        "Худшая игра 2025 года",
        "Лучший сюжет в играх 2025 года",
        "Лучшая инди игра 2025 года"
    ]
    with open(CATEGORIES_FILE, "w", encoding="utf-8") as f:
        json.dump(categories, f, ensure_ascii=False, indent=4)

# Загружаем категории при запуске
with open(CATEGORIES_FILE, "r", encoding="utf-8") as f:
    categories = json.load(f)

# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---
def is_english(text):
    return bool(re.match(r'^[A-Za-z0-9 !@#$%^&*(),.?":{}|<>-]+$', text))

def main_menu():
    """Главное меню"""
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add("Голосовать", "Создать голосование")
    markup.add("Предложить игру", "Предложить категорию")
    markup.add("Удалить", "Экспорт данных")
    return markup

def category_menu():
    """Меню выбора категорий с кнопкой 'Назад'"""
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    for cat in categories:
        markup.add(cat)
    # 🌟 Добавлена кнопка "Назад"
    markup.add("↩️ Назад в Главное меню") 
    return markup

# --- ОБРАБОТЧИКИ СООБЩЕНИЙ ---

@bot.message_handler(commands=["start"])
def start(message):
    bot.send_message(message.chat.id, "Привет! Я REPTILOID_GAME_AWARDS_BOT 🦖", reply_markup=main_menu())

@bot.message_handler(func=lambda m: True)
def menu_handler(message):
    if message.text == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Вы вернулись в главное меню.", reply_markup=main_menu())
    elif message.text == "Предложить категорию":
        msg = bot.send_message(message.chat.id, "Введите название новой категории (на русском):")
        bot.register_next_step_handler(msg, add_category)
    elif message.text == "Предложить игру":
        msg = bot.send_message(message.chat.id, "Выберите категорию для игры:", reply_markup=category_menu())
        bot.register_next_step_handler(msg, ask_game_name)
    elif message.text == "Голосовать":
        msg = bot.send_message(message.chat.id, "Выберите категорию для голосования:", reply_markup=category_menu())
        bot.register_next_step_handler(msg, show_games_for_vote)
    # 🌟 Новая функция "Создать голосование"
    elif message.text == "Создать голосование":
        msg = bot.send_message(message.chat.id, "Выберите категорию для создания опроса:", reply_markup=category_menu())
        bot.register_next_step_handler(msg, create_telegram_poll)
    # 🌟 Новая функция "Удалить"
    elif message.text == "Удалить":
        deletion_menu(message)
    elif message.text == "Экспорт данных":
        export_data(message)
    else:
        bot.send_message(message.chat.id, "Выберите действие с клавиатуры.", reply_markup=main_menu())

def add_category(message):
    new_cat = message.text.strip()
    if new_cat == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    global categories # Необходимо для изменения глобальной переменной
    if new_cat not in categories:
        categories.append(new_cat)
        with open(CATEGORIES_FILE, "w", encoding="utf-8") as f:
            json.dump(categories, f, ensure_ascii=False, indent=4)
        bot.send_message(message.chat.id, f"Категория '{new_cat}' добавлена!", reply_markup=main_menu())
    else:
        bot.send_message(message.chat.id, "Такая категория уже есть.", reply_markup=main_menu())

def ask_game_name(message):
    category = message.text
    if category == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    if category not in categories:
        bot.send_message(message.chat.id, "Неверная категория.", reply_markup=main_menu())
        return
    msg = bot.send_message(message.chat.id, f"Введите название игры на английском для категории '{category}':")
    bot.register_next_step_handler(msg, save_game, category)

def save_game(message, category):
    game = message.text.strip()
    if game == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    if not is_english(game):
        bot.send_message(message.chat.id, "Название должно быть только на английском.", reply_markup=main_menu())
        return

    df = pd.read_excel(DATA_FILE)
    if ((df["game"] == game) & (df["category"] == category)).any():
        bot.send_message(message.chat.id, "Эта игра уже предложена.", reply_markup=main_menu())
        return

    new_row = pd.DataFrame([{
        "user_id": message.from_user.id,
        "username": message.from_user.username or message.from_user.first_name,
        "category": category,
        "game": game,
        "votes": "",
        "poll_id": "" # Добавляем пустую колонку poll_id
    }])
    df = pd.concat([df, new_row], ignore_index=True)
    df.to_excel(DATA_FILE, index=False)
    bot.send_message(message.chat.id, f"Игра '{game}' добавлена в категорию '{category}'!", reply_markup=main_menu())

# --- 🌟 ЛОГИКА ГОЛОСОВАНИЯ (ИНЛАЙН КНОПКИ) ---

def show_games_for_vote(message):
    category = message.text
    if category == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    df = pd.read_excel(DATA_FILE)
    games = df[df["category"] == category]["game"].tolist()
    
    if not games:
        bot.send_message(message.chat.id, "В этой категории ещё нет игр.", reply_markup=main_menu())
        return

    markup = types.InlineKeyboardMarkup()
    for g in games:
        markup.add(types.InlineKeyboardButton(g, callback_data=f"vote|{category}|{g}"))

    bot.send_message(message.chat.id, "Нажмите на игру, за которую хотите проголосовать:", reply_markup=markup)

@bot.callback_query_handler(func=lambda call: call.data.startswith("vote|"))
def handle_vote(call):
    _, category, game = call.data.split("|")
    df = pd.read_excel(DATA_FILE)
    
    try:
        idx = df[(df["game"] == game) & (df["category"] == category)].index[0]
    except IndexError:
        bot.answer_callback_query(call.id, "Ошибка: Игра не найдена.")
        return

    # Проверка, голосовал ли уже пользователь
    votes = df.at[idx, "votes"]
    votes_list = votes.split(",") if pd.notna(votes) and votes else []
    user_tag = str(call.from_user.id)
    
    if user_tag in votes_list:
        bot.answer_callback_query(call.id, "Вы уже голосовали за эту игру!")
        return
        
    votes_list.append(user_tag)
    df.at[idx, "votes"] = ",".join(votes_list)
    df.to_excel(DATA_FILE, index=False)
    
    bot.answer_callback_query(call.id, f"Вы проголосовали за '{game}'!")


# --- 🌟 ЛОГИКА СОЗДАНИЯ NATIVE TELEGRAM POLL ---

def create_telegram_poll(message):
    category = message.text
    if category == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    df = pd.read_excel(DATA_FILE)
    category_df = df[df["category"] == category]
    games = category_df["game"].tolist()
    
    if len(games) < 2:
        bot.send_message(message.chat.id, "Нужно минимум две игры для создания опроса.", reply_markup=main_menu())
        return

    # Проверяем, был ли уже создан опрос для этой категории
    if category_df["poll_id"].astype(str).str.len().max() > 0:
        bot.send_message(message.chat.id, "Опрос для этой категории уже создан. Сначала удалите старый опрос.", reply_markup=main_menu())
        return

    # Создаем опрос (poll)
    try:
        # poll_options не должен превышать 100 символов, но игры на английском, так что ОК
        poll_message = bot.send_poll(
            chat_id=message.chat.id, 
            question=f"🏆 Категория: {category}",
            options=games, 
            is_anonymous=False, # Видим, кто за что голосует
            type='regular' # Обычный опрос, не викторина
        )
        
        # Сохраняем ID опроса, чтобы знать, какой опрос соответствует категории
        poll_id = poll_message.poll.id
        
        # Обновляем DataFrame: добавляем poll_id ко всем играм этой категории
        df.loc[df["category"] == category, "poll_id"] = poll_id
        df.to_excel(DATA_FILE, index=False)
        
        bot.send_message(message.chat.id, f"✅ Опрос для категории '{category}' успешно создан! ID опроса сохранен.", reply_markup=main_menu())

    except Exception as e:
        bot.send_message(message.chat.id, f"❌ Ошибка при создании опроса (возможно, не чат/группа): {e}", reply_markup=main_menu())


# --- 🌟 ЛОГИКА УДАЛЕНИЯ ---

def deletion_menu(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    markup.add("Удалить Категорию", "Удалить Игру")
    markup.add("Удалить Опрос", "↩️ Назад в Главное меню")
    bot.send_message(message.chat.id, "Что вы хотите удалить?", reply_markup=markup)
    bot.register_next_step_handler(message, handle_deletion_choice)

def handle_deletion_choice(message):
    choice = message.text
    if choice == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
    elif choice == "Удалить Категорию":
        msg = bot.send_message(message.chat.id, "Выберите категорию для удаления:", reply_markup=category_menu())
        bot.register_next_step_handler(msg, delete_category)
    elif choice == "Удалить Игру":
        msg = bot.send_message(message.chat.id, "Выберите категорию, из которой хотите удалить игру:", reply_markup=category_menu())
        bot.register_next_step_handler(msg, choose_game_to_delete)
    elif choice == "Удалить Опрос":
        msg = bot.send_message(message.chat.id, "Выберите категорию, опрос которой нужно удалить (сбросить ID):", reply_markup=category_menu())
        bot.register_next_step_handler(msg, delete_poll_id)
    else:
        bot.send_message(message.chat.id, "Неверный выбор.", reply_markup=main_menu())

def delete_category(message):
    cat_to_delete = message.text
    if cat_to_delete == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    global categories
    if cat_to_delete in categories:
        categories.remove(cat_to_delete)
        with open(CATEGORIES_FILE, "w", encoding="utf-8") as f:
            json.dump(categories, f, ensure_ascii=False, indent=4)
            
        # Удаляем все игры, связанные с этой категорией
        df = pd.read_excel(DATA_FILE)
        df = df[df["category"] != cat_to_delete]
        df.to_excel(DATA_FILE, index=False)
        
        bot.send_message(message.chat.id, f"Категория '{cat_to_delete}' и все связанные игры удалены.", reply_markup=main_menu())
    else:
        bot.send_message(message.chat.id, "Категория не найдена.", reply_markup=main_menu())

def choose_game_to_delete(message):
    category = message.text
    if category == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    if category not in categories:
        bot.send_message(message.chat.id, "Неверная категория.", reply_markup=main_menu())
        return
        
    df = pd.read_excel(DATA_FILE)
    games = df[df["category"] == category]["game"].tolist()
    
    if not games:
        bot.send_message(message.chat.id, "В этой категории нет игр для удаления.", reply_markup=main_menu())
        return
        
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    for g in games:
        markup.add(g)
    markup.add("↩️ Назад в Главное меню")
    
    msg = bot.send_message(message.chat.id, f"Выберите игру для удаления из категории '{category}':", reply_markup=markup)
    bot.register_next_step_handler(msg, delete_game, category)

def delete_game(message, category):
    game_to_delete = message.text
    if game_to_delete == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    df = pd.read_excel(DATA_FILE)
    # Ищем индекс строки с игрой и категорией
    idx_list = df[(df["game"] == game_to_delete) & (df["category"] == category)].index
    
    if not idx_list.empty:
        df = df.drop(idx_list)
        df.to_excel(DATA_FILE, index=False)
        bot.send_message(message.chat.id, f"Игра '{game_to_delete}' из категории '{category}' удалена.", reply_markup=main_menu())
    else:
        bot.send_message(message.chat.id, "Игра не найдена.", reply_markup=main_menu())

def delete_poll_id(message):
    category = message.text
    if category == "↩️ Назад в Главное меню":
        bot.send_message(message.chat.id, "Отмена операции.", reply_markup=main_menu())
        return
        
    df = pd.read_excel(DATA_FILE)
    
    # Сбрасываем poll_id для всех игр в выбранной категории
    # Используем fillna('') для работы с отсутствующими значениями
    if df[df["category"] == category]["poll_id"].astype(str).str.len().max() > 0:
        df.loc[df["category"] == category, "poll_id"] = ""
        df.to_excel(DATA_FILE, index=False)
        bot.send_message(message.chat.id, f"ID опроса для категории '{category}' сброшен. Вы можете создать новый опрос.", reply_markup=main_menu())
    else:
        bot.send_message(message.chat.id, f"Опрос для категории '{category}' не был создан.", reply_markup=main_menu())


# --- 📊 ЭКСПОРТ ДАННЫХ ---

def export_data(message):
    df = pd.read_excel(DATA_FILE)
    
    # XLSX
    df.to_excel("export.xlsx", index=False)
    
    # DOCX
    doc = Document()
    doc.add_heading("REPTILOID GAME AWARDS DATA", 0)
    for category in df["category"].unique():
        doc.add_heading(category, level=1)
        
        # Считаем количество голосов
        def count_votes(votes_str):
            if pd.isna(votes_str) or votes_str == "":
                return 0
            return len(str(votes_str).split(','))
            
        category_df = df[df["category"] == category].copy()
        category_df['vote_count'] = category_df['votes'].apply(count_votes)
        category_df = category_df.sort_values(by="vote_count", ascending=False)
        
        for _, row in category_df.iterrows():
            doc.add_paragraph(f"Игра: {row['game']} | Предложена: {row['username']} | Голосов: {row['vote_count']}")
            
    doc.save("export.docx")
    
    bot.send_message(message.chat.id, "Данные экспортированы: export.xlsx и export.docx.", reply_markup=main_menu())

# --- 🚨 ФУНКЦИЯ ДЛЯ ВЕБХУКА (ОБЯЗАТЕЛЬНА ДЛЯ PYTHONANYWHERE) 🚨 ---
def process_new_updates(update):
    """
    Обязательная функция для передачи обновлений от Flask в telebot.
    """
    bot.process_new_updates([update])